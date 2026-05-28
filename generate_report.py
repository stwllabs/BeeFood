from docx import Document
from docx.shared import Pt

doc = Document()

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)

# Cover / Title not requested directly, but include report details
heading('Laporan Proyek BeeFood', level=1)
para('Dokumen ini merupakan laporan proyek pengembangan aplikasi BeeFood untuk sistem pre-order makanan kampus.')

heading('DAFTAR ISI', level=2)
para('i. DAFTAR ISI')
para('ii. DAFTAR LAMPIRAN')
para('1. BAB 1. PENDAHULUAN')
para('1.1 Latar Belakang')
para('1.2 Tujuan')
para('1.3 Prediksi Manfaat')
para('1.4 Luaran')
para('2. BAB 2. TINJAUAN PUSTAKA')
para('2.1 Sistem Pre-Order Makanan dan Pengurangan Antrian')
para('2.2 Teknologi Frontend dan Backend untuk Aplikasi Kampus')
para('3. BAB 3. TAHAP PELAKSANAAN')
para('3.1 Deskripsi Produk/Alat/Sistem')
para('3.2 Alur dan Tahapan Pelaksanaan')
para('3.3 Perancangan Produk/Alat/Sistem')
para('3.4 Pengujian')
para('4. BAB 4. BIAYA DAN JADWAL KEGIATAN')
para('4.1 Anggaran Biaya')
para('4.2 Jadwal Kegiatan')
para('DAFTAR PUSTAKA')
para('LAMPIRAN')

heading('DAFTAR GAMBAR', level=2)
para('Tidak ada gambar dalam laporan ini (opsional).')

heading('DAFTAR TABEL', level=2)
para('Tidak ada tabel khusus dalam laporan ini (opsional).')

heading('DAFTAR LAMPIRAN', level=2)
para('Lampiran 1. Biodata Ketua dan Anggota, serta Dosen Pendamping')
para('Lampiran 2. Justifikasi Anggaran Kegiatan')
para('Lampiran 3. Susunan Tim Pengusul dan Pembagian Tugas')
para('Lampiran 4. Surat Pernyataan Ketua Pengusul')
para('Lampiran 5. Gambaran Teknologi yang akan Dikembangkan')
para('Lampiran 6. Hasil Uji Periksa Similaritas Proposal')
para('Lampiran 7. Dokumentasi Teknis dan Detail Pengembangan (opsional)')

heading('BAB 1. PENDAHULUAN', level=1)
heading('Latar Belakang', level=2)
para('Proyek ini mengembangkan aplikasi BeeFood, sebuah platform pre-order makanan kampus yang ditujukan untuk mahasiswa dan tenant kantin.')
para('Masalah yang dihadapi di lingkungan kampus adalah antrean panjang, waktu tunggu di kantin, dan manajemen pesanan yang kurang efisien.')
para('Dengan memanfaatkan teknologi web modern, aplikasi ini mempertemukan mahasiswa sebagai pemesan, tenant kantin sebagai penjual, dan backend untuk autentikasi, manajemen menu, pre-order, dan update status.')

heading('Tujuan', level=2)
para('Tujuan proyek BeeFood adalah membuat aplikasi web untuk pre-order makanan kampus, menyediakan dashboard khusus untuk mahasiswa dan tenant, mengurangi antrean fisik, serta mempercepat proses pemesanan dan pelacakan status pesanan.')

heading('Prediksi Manfaat', level=2)
para('Proyek ini diharapkan memberikan manfaat bagi mahasiswa dalam memesan makanan dari kelas dengan notifikasi status, bagi tenant dalam manajemen antrean dan status pesanan, serta bagi kampus dalam mengurangi kerumunan kantin dan meningkatkan kepuasan layanan.')

heading('Luaran', level=2)
para('Luaran yang dihasilkan meliputi aplikasi frontend React + Vite, backend Express + Prisma, database PostgreSQL, serta dokumentasi teknis arsitektur, alur, dan pengujian.')

heading('BAB 2. TINJAUAN PUSTAKA', level=1)
heading('Sistem Pre-Order Makanan dan Pengurangan Antrian', level=2)
para('Topik yang relevan meliputi sistem pemesanan makanan online untuk lingkungan kampus, studi tentang dampak pre-order terhadap waktu tunggu, dan konsep e-wallet kampus untuk pembelian makanan digital.')

heading('Teknologi Frontend dan Backend untuk Aplikasi Kampus', level=2)
para('Topik yang relevan meliputi penggunaan React + Vite, Express.js, Prisma, PostgreSQL, Socket.io, serta desain UI/UX form login, dashboard, dan antarmuka pemesanan.')

heading('BAB 3. TAHAP PELAKSANAAN', level=1)
heading('Deskripsi Produk/Alat/Sistem', level=2)
para('BeeFood adalah sistem aplikasi pre-order makanan kampus yang terdiri dari frontend React, backend Express, serta database PostgreSQL dengan model User, Tenant, Menu, dan Order.')
para('Fitur utama termasuk login role student/tenant, dashboard mahasiswa, dashboard tenant, pre-order, checkout, dan update status pesanan.')

heading('Alur dan Tahapan Pelaksanaan', level=2)
para('Alur sistem dimulai dari login, pemilihan role, pemesanan menu, checkout, sampai tenant memproses dan memperbarui status order secara realtime.')
para('Tahapan pelaksanaan mencakup analisis kebutuhan, desain UI/UX, implementasi frontend dan backend, integrasi API, pengujian, dan dokumentasi.')

heading('Perancangan Produk/Alat/Sistem', level=2)
para('Perancangan meliputi struktur folder frontend, routing React, desain basis data Prisma, serta API backend untuk autentikasi, menu, order, dan status update.')

heading('Pengujian', level=2)
para('Pengujian meliputi fungsional frontend, pengujian backend API, integrasi antara frontend dan backend, serta pengujian alur user mahasiswa dan tenant.')

heading('BAB 4. BIAYA DAN JADWAL KEGIATAN', level=1)
heading('Anggaran Biaya', level=2)
para('Estimasi biaya proyek contoh: server dan hosting Rp 1.000.000, domain/cloud deployment Rp 500.000, biaya pengembangan tim Rp 2.500.000, total estimasi Rp 4.000.000.')

heading('Jadwal Kegiatan', level=2)
para('Contoh jadwal 4 minggu: Minggu 1 analisis kebutuhan dan perancangan, Minggu 2 implementasi frontend dan backend, Minggu 3 integrasi API dan manajemen order, Minggu 4 pengujian, perbaikan bug, dokumentasi, dan presentasi.')

heading('DAFTAR PUSTAKA', level=1)
para('Dokumentasi React, Vite, React Router, Express.js, Socket.io, Prisma, PostgreSQL, dan referensi studi sistem pre-order makanan kampus.')

heading('LAMPIRAN', level=1)
heading('Lampiran 1. Biodata Ketua dan Anggota, serta Dosen Pendamping', level=2)
para('Isi biodata ketua, anggota, dan dosen pendamping sesuai tim proyek.')

heading('Lampiran 2. Justifikasi Anggaran Kegiatan', level=2)
para('Penjelasan penggunaan biaya hosting, deployment, dan pengembangan.')

heading('Lampiran 3. Susunan Tim Pengusul dan Pembagian Tugas', level=2)
para('Pembagian tugas antara pengembang frontend, backend, dokumentasi, dan pengujian.')

heading('Lampiran 4. Surat Pernyataan Ketua Pengusul', level=2)
para('Pernyataan keaslian dan komitmen penyelesaian proyek.')

heading('Lampiran 5. Gambaran Teknologi yang akan Dikembangkan', level=2)
para('Detail teknologi yang digunakan: React, Vite, Express, Prisma, Socket.io, PostgreSQL.')

heading('Lampiran 6. Hasil Uji Periksa Similaritas Proposal', level=2)
para('Laporan pemeriksaan plagiarisme jika ada.')

heading('Lampiran 7. Dokumentasi Teknis dan Detail Pengembangan', level=2)
para('Dokumentasi tambahan berupa diagram alir, skrinsut antarmuka, atau penjelasan kode.')

output_path = 'Laporan BeeFood.docx'
doc.save(output_path)
print(f'Created {output_path}')
